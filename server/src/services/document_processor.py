"""
Processadores de documentos para a base de conhecimento.
Extrai texto de diferentes tipos de arquivo e prepara para vetorização.
"""

import os
import aiofiles
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging
from abc import ABC, abstractmethod

# Imports para processamento de arquivos
try:
    import PyPDF2
    from docx import Document
    import markdown
except ImportError as e:
    logging.warning(f"Algumas bibliotecas de processamento não estão disponíveis: {e}")

from ..core.config import settings

logger = logging.getLogger(__name__)

class DocumentProcessor(ABC):
    """Classe base abstrata para processadores de documentos"""

    @abstractmethod
    async def extract_text(self, file_path: Path) -> str:
        """Extrai texto do arquivo"""
        pass

    @abstractmethod
    def get_supported_extensions(self) -> List[str]:
        """Retorna extensões suportadas"""
        pass

    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """
        Divide texto em chunks para processamento eficiente

        Args:
            text: Texto a ser dividido
            chunk_size: Tamanho máximo do chunk
            overlap: Sobreposição entre chunks

        Returns:
            List[str]: Lista de chunks de texto
        """
        if chunk_size is None:
            chunk_size = settings.CHUNK_SIZE
        if overlap is None:
            overlap = settings.CHUNK_OVERLAP

        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # Tentar quebrar em uma sentença completa
            if end < len(text):
                # Procurar por ponto final próximo
                for i in range(end, max(start + chunk_size - 100, start), -1):
                    if text[i] in '.!?\\n':
                        end = i + 1
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - overlap

            # Evitar loop infinito
            if start >= end:
                start = end

        return chunks

class PDFProcessor(DocumentProcessor):
    """Processador para arquivos PDF"""

    async def extract_text(self, file_path: Path) -> str:
        """Extrai texto de arquivo PDF"""
        try:
            text_content = []

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(f"--- Página {page_num + 1} ---\\n{text}")
                    except Exception as e:
                        logger.warning(f"Erro ao extrair página {page_num + 1}: {e}")
                        continue

            full_text = "\\n".join(text_content)
            logger.info(f"✅ PDF processado: {len(full_text)} caracteres extraídos")

            return full_text

        except Exception as e:
            logger.error(f"❌ Erro ao processar PDF {file_path}: {e}")
            raise

    def get_supported_extensions(self) -> List[str]:
        return ['pdf']

class DOCXProcessor(DocumentProcessor):
    """Processador para arquivos DOCX"""

    async def extract_text(self, file_path: Path) -> str:
        """Extrai texto de arquivo DOCX"""
        try:
            doc = Document(file_path)
            paragraphs = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            # Extrair texto de tabelas também
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        paragraphs.append(row_text)

            full_text = "\\n".join(paragraphs)
            logger.info(f"✅ DOCX processado: {len(full_text)} caracteres extraídos")

            return full_text

        except Exception as e:
            logger.error(f"❌ Erro ao processar DOCX {file_path}: {e}")
            raise

    def get_supported_extensions(self) -> List[str]:
        return ['docx']

class TXTProcessor(DocumentProcessor):
    """Processador para arquivos de texto simples"""

    async def extract_text(self, file_path: Path) -> str:
        """Extrai texto de arquivo TXT"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()

            logger.info(f"✅ TXT processado: {len(content)} caracteres extraídos")
            return content

        except UnicodeDecodeError:
            # Tentar com encoding diferente
            try:
                async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
                    content = await file.read()
                logger.info(f"✅ TXT processado (latin-1): {len(content)} caracteres")
                return content
            except Exception as e:
                logger.error(f"❌ Erro ao processar TXT {file_path}: {e}")
                raise
        except Exception as e:
            logger.error(f"❌ Erro ao processar TXT {file_path}: {e}")
            raise

    def get_supported_extensions(self) -> List[str]:
        return ['txt']

class MarkdownProcessor(DocumentProcessor):
    """Processador para arquivos Markdown"""

    async def extract_text(self, file_path: Path) -> str:
        """Extrai texto de arquivo Markdown"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                md_content = await file.read()

            # Converter Markdown para texto simples
            html = markdown.markdown(md_content)

            # Remover tags HTML básicas (implementação simples)
            import re
            text = re.sub(r'<[^>]+>', '', html)
            text = text.replace('&nbsp;', ' ').replace('&amp;', '&')

            logger.info(f"✅ Markdown processado: {len(text)} caracteres extraídos")
            return text

        except Exception as e:
            logger.error(f"❌ Erro ao processar Markdown {file_path}: {e}")
            raise

    def get_supported_extensions(self) -> List[str]:
        return ['md', 'markdown']

class DocumentProcessorService:
    """
    Serviço principal para processamento de documentos

    Gerencia diferentes tipos de processadores e coordena a extração
    e preparação de texto para o banco vetorial.
    """

    def __init__(self):
        self.processors = {
            'pdf': PDFProcessor(),
            'docx': DOCXProcessor(),
            'txt': TXTProcessor(),
            'md': MarkdownProcessor()
        }

    def get_processor_for_file(self, file_path: Path) -> Optional[DocumentProcessor]:
        """Retorna o processador apropriado para um arquivo"""
        extension = file_path.suffix.lower().lstrip('.')

        for processor in self.processors.values():
            if extension in processor.get_supported_extensions():
                return processor

        return None

    async def process_document(
        self,
        file_path: Path,
        metadata: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """
        Processa um documento completamente

        Args:
            file_path: Caminho para o arquivo
            metadata: Metadados adicionais

        Returns:
            Dict com texto extraído, chunks e metadados
        """
        try:
            # Validar arquivo
            if not file_path.exists():
                raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")

            # Obter processador
            processor = self.get_processor_for_file(file_path)
            if not processor:
                raise ValueError(f"Tipo de arquivo não suportado: {file_path.suffix}")

            # Extrair texto
            logger.info(f"🔄 Processando documento: {file_path.name}")
            extracted_text = await processor.extract_text(file_path)

            if not extracted_text.strip():
                raise ValueError("Nenhum texto extraído do documento")

            # Criar chunks
            chunks = processor.chunk_text(extracted_text)

            # Preparar metadados
            file_metadata = {
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'file_type': file_path.suffix.lower().lstrip('.'),
                'word_count': len(extracted_text.split()),
                'chunk_count': len(chunks),
                'processed_at': str(Path.ctime(file_path)),
                **(metadata or {})
            }

            result = {
                'text': extracted_text,
                'chunks': chunks,
                'metadata': file_metadata,
                'success': True
            }

            logger.info(
                f"✅ Documento processado: {len(chunks)} chunks, "
                f"{file_metadata['word_count']} palavras"
            )

            return result

        except Exception as e:
            logger.error(f"❌ Erro ao processar documento {file_path}: {e}")
            return {
                'text': '',
                'chunks': [],
                'metadata': metadata or {},
                'success': False,
                'error': str(e)
            }

    def get_supported_extensions(self) -> List[str]:
        """Retorna todas as extensões suportadas"""
        extensions = []
        for processor in self.processors.values():
            extensions.extend(processor.get_supported_extensions())
        return list(set(extensions))

    async def validate_file(self, file_path: Path) -> Dict[str, Any]:
        """
        Valida se um arquivo pode ser processado

        Returns:
            Dict com status de validação
        """
        validation = {
            'is_valid': False,
            'errors': [],
            'warnings': []
        }

        # Verificar se existe
        if not file_path.exists():
            validation['errors'].append("Arquivo não encontrado")
            return validation

        # Verificar extensão
        extension = file_path.suffix.lower().lstrip('.')
        if extension not in self.get_supported_extensions():
            validation['errors'].append(f"Extensão '{extension}' não suportada")
            return validation

        # Verificar tamanho
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > settings.MAX_FILE_SIZE_MB:
            validation['errors'].append(
                f"Arquivo muito grande: {file_size_mb:.1f}MB "
                f"(máximo: {settings.MAX_FILE_SIZE_MB}MB)"
            )
            return validation

        # Verificar se pode ser lido
        try:
            processor = self.get_processor_for_file(file_path)
            if not processor:
                validation['errors'].append("Processador não encontrado")
                return validation
        except Exception as e:
            validation['errors'].append(f"Erro ao acessar processador: {e}")
            return validation

        # Se chegou até aqui, é válido
        validation['is_valid'] = True

        # Adicionar avisos se necessário
        if file_size_mb > 5:
            validation['warnings'].append("Arquivo grande pode demorar para processar")

        return validation

# Instância global
document_processor = DocumentProcessorService()